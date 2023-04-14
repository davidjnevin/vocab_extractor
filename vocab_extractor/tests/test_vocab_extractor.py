import sqlite3
import pytest
from docx import Document
from pathlib import Path

from vocab_extractor.app import (
    extract_vocabulary,
    read_file,
    create_database,
    insert_word,
    process_reports,
    batch_files,
    process_all_reports,
)


def test_extract_vocabulary():
    text = "This is a sample text. It contains some words, punctuation, and numbers like 1234."
    expected_result = {"This", "is", "a", "sample", "text", "It", "contains", "some", "words", "punctuation", "and", "numbers", "like", "1234"}
    assert extract_vocabulary(text) == expected_result


def test_read_file(tmp_path: Path):
    # Create a temporary .docx file for testing
    temp_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a temporary file for testing.")
    doc.save(temp_docx)

    content = read_file(str(temp_docx))
    assert content.strip() == "This is a temporary file for testing."


def test_create_database(tmp_path: Path):
    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    assert isinstance(conn, sqlite3.Connection)


def test_insert_word(tmp_path: Path):
    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    word = "test"
    insert_word(word, conn)
    cursor = conn.execute("SELECT word FROM vocabulary WHERE word=?", (word,))
    assert cursor.fetchone()[0] == word


def test_process_reports(tmp_path: Path):
    # Create a temporary .docx file for testing
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document for process_reports.")
    doc.save(test_docx)

    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    report_files = [str(test_docx)]
    process_reports(report_files, conn)
    cursor = conn.execute("SELECT word FROM vocabulary")
    words = [row[0] for row in cursor.fetchall()]
    assert "test" in words


def test_batch_files():
    report_files = list(range(10))
    batch_gen = batch_files(report_files, batch_size=3)
    batches = [batch for batch in batch_gen]
    assert len(batches) == 4


def test_process_all_reports(tmp_path: Path):
    # Create a temporary .docx file for testing
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document for process_all_reports.")
    doc.save(test_docx)

    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    report_files = [str(test_docx)]
    process_all_reports(report_files, conn)
    cursor = conn.execute("SELECT word FROM vocabulary")
    words = [row[0] for row in cursor.fetchall()]
    assert "test" in words

