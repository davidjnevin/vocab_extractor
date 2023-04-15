import sqlite3
from pathlib import Path
import pytest

from docx import Document

from vocab_extractor.app import (
    process_reports,
    batch_files,
    process_all_reports,
)

from vocab_extractor.io_utils import (
    create_database, read_file
)


def test_create_database(tmp_path: Path):
    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    assert isinstance(conn, sqlite3.Connection)


def test_process_reports(tmp_path: Path):
    # Create a temporary .docx file for testing
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document for process_reports.")
    doc.save(test_docx)

    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    report_files = [str(test_docx)]
    process_reports(report_files, conn)
    with conn:
        cursor = conn.execute("SELECT word, pos FROM vocabulary")
        words = [(row[0], row[1]) for row in cursor.fetchall()]
    assert ("test", "NN") in words


def test_batch_files():
    report_files = list(range(10))
    batch_gen = batch_files(report_files, batch_size=3)
    batches = [batch for batch in batch_gen]
    assert len(batches) == 4


def test_process_all_reports(tmp_path: Path):
    # Create a temporary .docx file for testing
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document for process_all_reports.")
    doc.save(test_docx)

    temp_db = tmp_path / "test.db"
    conn = create_database(str(temp_db))
    report_files = [str(test_docx)]
    process_all_reports(report_files, conn)
    with conn:
        cursor = conn.execute("SELECT word, pos FROM vocabulary")
        words = [(row[0], row[1]) for row in cursor.fetchall()]
    assert ("test", "NN") in words

